"""
DOCX Exporter Module
Exports policy packages to Microsoft Word format
"""

import re
from pathlib import Path
from typing import List, Optional, Dict, Any
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .package_builder import PackageResult, PolicyDocument


class DocxExporter:
    """
    Exports policy packages to DOCX format.

    Usage:
        exporter = DocxExporter()
        exporter.export_package(result, "output/acme_policies.docx")
    """

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX export. "
                "Install it with: pip install python-docx"
            )

    def create_document(self, client_name: str) -> Document:
        """Create a new Word document with basic styles"""
        doc = Document()

        # Set up styles
        self._setup_styles(doc)

        return doc

    def _setup_styles(self, doc: Document):
        """Set up document styles"""
        styles = doc.styles

        # Modify Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)

        # Modify Heading styles
        for i in range(1, 5):
            style_name = f'Heading {i}'
            if style_name in styles:
                heading_style = styles[style_name]
                heading_font = heading_style.font
                heading_font.name = 'Calibri'
                heading_font.bold = True
                if i == 1:
                    heading_font.size = Pt(18)
                    heading_font.color.rgb = RGBColor(0, 51, 102)
                elif i == 2:
                    heading_font.size = Pt(14)
                    heading_font.color.rgb = RGBColor(0, 51, 102)
                elif i == 3:
                    heading_font.size = Pt(12)
                else:
                    heading_font.size = Pt(11)

    def _parse_markdown_to_docx(self, doc: Document, content: str):
        """Convert markdown content to Word document elements"""
        lines = content.split('\n')
        in_list = False
        list_items = []

        for line in lines:
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                if in_list and list_items:
                    self._add_list(doc, list_items)
                    list_items = []
                    in_list = False
                continue

            # Headings
            if stripped.startswith('######'):
                doc.add_heading(stripped[6:].strip(), level=6)
            elif stripped.startswith('#####'):
                doc.add_heading(stripped[5:].strip(), level=5)
            elif stripped.startswith('####'):
                doc.add_heading(stripped[4:].strip(), level=4)
            elif stripped.startswith('###'):
                doc.add_heading(stripped[3:].strip(), level=3)
            elif stripped.startswith('##'):
                doc.add_heading(stripped[2:].strip(), level=2)
            elif stripped.startswith('#'):
                doc.add_heading(stripped[1:].strip(), level=1)

            # Lists
            elif stripped.startswith('- ') or stripped.startswith('* '):
                in_list = True
                list_items.append(stripped[2:])

            elif re.match(r'^\d+\.\s', stripped):
                in_list = True
                list_items.append(re.sub(r'^\d+\.\s', '', stripped))

            # Regular paragraphs
            else:
                if in_list and list_items:
                    self._add_list(doc, list_items)
                    list_items = []
                    in_list = False

                # Process inline formatting
                formatted_text = self._process_inline_formatting(stripped)
                p = doc.add_paragraph()
                self._add_formatted_runs(p, formatted_text)

        # Handle remaining list items
        if list_items:
            self._add_list(doc, list_items)

    def _add_list(self, doc: Document, items: List[str]):
        """Add a bulleted list to the document"""
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            formatted_text = self._process_inline_formatting(item)
            self._add_formatted_runs(p, formatted_text)

    def _process_inline_formatting(self, text: str) -> List[Dict[str, Any]]:
        """Process inline markdown formatting and return list of text runs"""
        runs = []
        current_pos = 0

        # Find bold (**text**) and italic (*text*) markers
        pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'

        for match in re.finditer(pattern, text):
            # Add text before the match
            if match.start() > current_pos:
                runs.append({
                    'text': text[current_pos:match.start()],
                    'bold': False,
                    'italic': False,
                    'code': False
                })

            matched_text = match.group()
            if matched_text.startswith('**') and matched_text.endswith('**'):
                runs.append({
                    'text': matched_text[2:-2],
                    'bold': True,
                    'italic': False,
                    'code': False
                })
            elif matched_text.startswith('`') and matched_text.endswith('`'):
                runs.append({
                    'text': matched_text[1:-1],
                    'bold': False,
                    'italic': False,
                    'code': True
                })
            elif matched_text.startswith('*') and matched_text.endswith('*'):
                runs.append({
                    'text': matched_text[1:-1],
                    'bold': False,
                    'italic': True,
                    'code': False
                })

            current_pos = match.end()

        # Add remaining text
        if current_pos < len(text):
            runs.append({
                'text': text[current_pos:],
                'bold': False,
                'italic': False,
                'code': False
            })

        return runs if runs else [{'text': text, 'bold': False, 'italic': False, 'code': False}]

    def _add_formatted_runs(self, paragraph, runs: List[Dict[str, Any]]):
        """Add formatted text runs to a paragraph"""
        for run_data in runs:
            run = paragraph.add_run(run_data['text'])
            run.bold = run_data.get('bold', False)
            run.italic = run_data.get('italic', False)
            if run_data.get('code', False):
                run.font.name = 'Consolas'
                run.font.size = Pt(10)

    def export_policy(self, doc: Document, policy: PolicyDocument, include_metadata: bool = True):
        """Export a single policy to the document"""
        # Policy title
        doc.add_heading(policy.title, level=1)

        # Metadata section
        if include_metadata:
            meta_para = doc.add_paragraph()
            meta_para.add_run("Policy ID: ").bold = True
            meta_para.add_run(policy.id)
            meta_para.add_run(" | ")
            meta_para.add_run("Category: ").bold = True
            meta_para.add_run(policy.category.replace('-', ' ').title())

            if policy.frameworks:
                framework_text = ", ".join(policy.frameworks.keys())
                meta_para.add_run(" | ")
                meta_para.add_run("Frameworks: ").bold = True
                meta_para.add_run(framework_text.upper())

        # Add a line separator
        doc.add_paragraph("─" * 50)

        # Policy content
        self._parse_markdown_to_docx(doc, policy.content)

        # Add page break before next policy
        doc.add_page_break()

    def export_package(self, result: PackageResult, output_path: str,
                       include_toc: bool = True,
                       include_metadata: bool = True):
        """
        Export a complete policy package to DOCX.

        Args:
            result: PackageResult from PackageBuilder
            output_path: Path for the output DOCX file
            include_toc: Include table of contents
            include_metadata: Include policy metadata headers
        """
        doc = self.create_document(result.client_name)

        # Cover page
        self._add_cover_page(doc, result)

        # Table of Contents
        if include_toc:
            self._add_table_of_contents(doc, result)

        # Export each policy
        for policy in sorted(result.policies, key=lambda p: (p.category, p.title)):
            self.export_policy(doc, policy, include_metadata)

        # Save document
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))

        return output

    def _add_cover_page(self, doc: Document, result: PackageResult):
        """Add a cover page to the document"""
        # Add some spacing at top
        for _ in range(5):
            doc.add_paragraph()

        # Title
        title = doc.add_heading(f"{result.client_name}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading("Information Security Policy Package", level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing
        for _ in range(3):
            doc.add_paragraph()

        # Generation info
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info.add_run(f"Generated: {result.generated_at.strftime('%B %d, %Y')}\n")
        info.add_run(f"Total Policies: {result.total_policies}\n")

        if result.frameworks_covered:
            info.add_run(f"Frameworks: {', '.join(fw.upper() for fw in result.frameworks_covered)}\n")

        if result.incomplete_count > 0:
            info.add_run(f"\n⚠️ {result.incomplete_count} policies require customization")

        # Page break
        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document, result: PackageResult):
        """Add a table of contents"""
        doc.add_heading("Table of Contents", level=1)

        # Group policies by category
        categories = {}
        for policy in result.policies:
            cat = policy.category or "Uncategorized"
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(policy)

        for category in sorted(categories.keys()):
            # Category heading
            doc.add_heading(category.replace('-', ' ').title(), level=2)

            # Policy list
            for policy in sorted(categories[category], key=lambda p: p.title):
                status = "⚠️ " if policy.incomplete_sections else ""
                p = doc.add_paragraph(f"{status}{policy.title}", style='List Bullet')

        doc.add_page_break()

    def export_single_policy(self, policy: PolicyDocument, output_path: str,
                             client_name: str = ""):
        """Export a single policy to its own DOCX file"""
        doc = self.create_document(client_name)
        self.export_policy(doc, policy, include_metadata=True)

        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))

        return output


def main():
    """Test the DOCX exporter"""
    import sys
    from pathlib import Path

    # Add project to path
    script_dir = Path(__file__).parent
    project_root = script_dir.parent.parent
    sys.path.insert(0, str(project_root / "src"))

    from generation.package_builder import PackageBuilder, ClientConfig

    policies_dir = project_root / "policies"
    frameworks_dir = project_root / "config" / "frameworks"
    output_dir = project_root / "output"

    print("=" * 70)
    print("DOCX EXPORTER TEST")
    print("=" * 70)

    # Build a test package
    builder = PackageBuilder(str(policies_dir), str(frameworks_dir))

    config = ClientConfig(
        name="Acme Corporation",
        variables={
            "ORGANIZATION_NAME": "Acme Corporation",
            "CSO_TITLE": "Chief Information Security Officer",
            "EXEC_MGMT": "Executive Leadership Team",
            "IT_STAFF": "IT Department",
            "HR_DEPARTMENT": "Human Resources",
            "LEGAL_DEPARTMENT": "Legal Department",
            "RMO_TITLE": "Risk Management Officer"
        },
        frameworks=["soc2"]  # Just SOC 2 for test
    )

    print(f"\nBuilding package for: {config.name}")
    result = builder.build_package(config)
    print(f"Total policies: {result.total_policies}")

    # Export to DOCX
    try:
        exporter = DocxExporter()
        output_file = output_dir / f"{config.name.lower().replace(' ', '_')}_policies.docx"

        print(f"\nExporting to: {output_file}")
        exporter.export_package(result, str(output_file))
        print(f"✅ Export complete: {output_file}")

    except ImportError as e:
        print(f"\n❌ {e}")
        print("DOCX export requires python-docx. Install with: pip install python-docx")


if __name__ == "__main__":
    main()
