"""
Policy Parser Module
Converts DOCX policies to Markdown + YAML frontmatter format
Parses and validates policy structure
"""

import os
import re
import yaml
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional, List, Dict, Any, Set
from datetime import datetime
from enum import Enum

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import frontmatter
    FRONTMATTER_AVAILABLE = True
except ImportError:
    FRONTMATTER_AVAILABLE = False


class PolicyType(Enum):
    """Types of policy documents"""
    POLICY = "policy"
    PROCEDURE = "procedure"
    PLAN = "plan"
    AGREEMENT = "agreement"
    STANDARD = "standard"
    GUIDELINE = "guideline"


class PolicyStatus(Enum):
    """Status of a policy"""
    DRAFT = "draft"
    ACTIVE = "active"
    UNDER_REVIEW = "under_review"
    DEPRECATED = "deprecated"
    ARCHIVED = "archived"


@dataclass
class PolicySection:
    """Represents a section within a policy"""
    number: str  # e.g., "I", "II.A", "III.B.2"
    title: str
    content: str
    level: int  # Heading level (1-6)
    subsections: List['PolicySection'] = field(default_factory=list)


@dataclass
class IncompletenessMarker:
    """Marks a section requiring customization"""
    section: str
    reason: str
    frameworks: List[str]
    priority: str = "medium"  # low, medium, high, critical


@dataclass
class Policy:
    """Complete policy document structure"""
    # Identification
    id: str
    title: str
    filename: str

    # Classification
    category: str
    type: PolicyType
    status: PolicyStatus
    version: str = "1.0.0"

    # Content
    sections: List[PolicySection] = field(default_factory=list)
    raw_content: str = ""

    # Compliance mapping
    frameworks: Dict[str, List[str]] = field(default_factory=dict)

    # Cross-references
    references: List[str] = field(default_factory=list)

    # Variables used
    variables: List[str] = field(default_factory=list)

    # Conditions
    conditions: List[Dict[str, Any]] = field(default_factory=list)

    # Incompleteness
    requires_customization: List[IncompletenessMarker] = field(default_factory=list)

    # Organization tiers
    organization_tiers: List[str] = field(default_factory=list)

    # Industries
    industries: List[str] = field(default_factory=list)

    # Metadata
    created: Optional[str] = None
    last_reviewed: Optional[str] = None
    next_review: Optional[str] = None
    author: Optional[str] = None
    approver: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for YAML serialization"""
        return {
            'id': self.id,
            'title': self.title,
            'version': self.version,
            'category': self.category,
            'type': self.type.value,
            'status': self.status.value,
            'frameworks': self.frameworks,
            'references': self.references,
            'variables': self.variables,
            'conditions': self.conditions,
            'requires_customization': [
                {
                    'section': m.section,
                    'reason': m.reason,
                    'frameworks': m.frameworks,
                    'priority': m.priority
                }
                for m in self.requires_customization
            ],
            'organization_tiers': self.organization_tiers,
            'industries': self.industries,
            'created': self.created,
            'last_reviewed': self.last_reviewed,
            'next_review': self.next_review,
            'author': self.author,
            'approver': self.approver
        }

    def to_markdown(self) -> str:
        """Convert to Markdown with YAML frontmatter"""
        # Build frontmatter
        frontmatter_dict = self.to_dict()

        # Build markdown content
        content_lines = []
        for section in self.sections:
            content_lines.append(self._section_to_markdown(section))

        content = '\n\n'.join(content_lines)

        # Combine
        yaml_str = yaml.dump(frontmatter_dict, default_flow_style=False, sort_keys=False)
        return f"---\n{yaml_str}---\n\n{content}"

    def _section_to_markdown(self, section: PolicySection, depth: int = 0) -> str:
        """Convert a section to markdown"""
        lines = []

        # Heading
        heading_prefix = '#' * min(section.level + depth, 6)
        if section.number:
            lines.append(f"{heading_prefix} {section.number}. {section.title}")
        else:
            lines.append(f"{heading_prefix} {section.title}")

        # Content
        if section.content:
            lines.append("")
            lines.append(section.content)

        # Subsections
        for subsection in section.subsections:
            lines.append("")
            lines.append(self._section_to_markdown(subsection, depth + 1))

        return '\n'.join(lines)


class PolicyParser:
    """
    Parses policy documents from various formats
    Converts DOCX to structured Policy objects
    """

    # Common variables found in policies
    KNOWN_VARIABLES = [
        'ORGANIZATION_NAME',
        'EFFECTIVE_DATE',
        'APPROVAL_DATE',
        'APPROVER',
        'EXEC_MGMT',
        'CSO_TITLE',
        'VERSION',
        'IT_STAFF',
        'RMO_TITLE',
        'PRIVACY_OFFICER',
        'HR_DEPARTMENT',
        'LEGAL_DEPARTMENT',
        'REVIEW_DATE',
        'COMPANY_ADDRESS',
        'CONTACT_EMAIL'
    ]

    # Section number patterns
    ROMAN_NUMERALS = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']

    # Category keywords for auto-detection
    CATEGORY_KEYWORDS = {
        'access-control': ['access control', 'authentication', 'authorization', 'password', 'account'],
        'network-security': ['network', 'firewall', 'vpn', 'router', 'wireless', 'dns'],
        'data-protection': ['data classification', 'data privacy', 'encryption', 'data retention'],
        'business-continuity': ['business continuity', 'disaster recovery', 'incident response', 'backup'],
        'compliance': ['compliance', 'hipaa', 'pci', 'gdpr', 'soc', 'audit'],
        'operations': ['change management', 'configuration', 'patch', 'maintenance'],
        'physical-security': ['physical security', 'facility', 'physical access'],
        'vendor-management': ['vendor', 'third party', 'outsourcing', 'supplier'],
        'risk-management': ['risk assessment', 'risk management', 'vulnerability'],
        'personnel': ['personnel', 'training', 'awareness', 'staffing'],
        'development': ['software development', 'sdlc', 'application'],
        'endpoint': ['workstation', 'mobile device', 'byod', 'portable', 'smartphone'],
        'governance': ['governance', 'policy', 'framework', 'ethics']
    }

    def __init__(self, policies_dir: Optional[str] = None):
        self.policies_dir = policies_dir
        self._policies_cache: Dict[str, Policy] = {}

    def parse_docx(self, filepath: str) -> Policy:
        """
        Parse a DOCX file into a Policy object

        Args:
            filepath: Path to DOCX file

        Returns:
            Parsed Policy object
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

        doc = Document(filepath)
        filename = os.path.basename(filepath)
        title = os.path.splitext(filename)[0]

        # Generate ID from title
        policy_id = self._generate_id(title)

        # Detect policy type
        policy_type = self._detect_type(title)

        # Detect category
        category = self._detect_category(title, doc)

        # Extract sections
        sections = self._extract_sections(doc)

        # Extract raw content
        raw_content = self._extract_raw_content(doc)

        # Find variables used
        variables = self._find_variables(raw_content)

        # Find cross-references
        references = self._find_references(raw_content)

        # Create policy
        policy = Policy(
            id=policy_id,
            title=title,
            filename=filename,
            category=category,
            type=policy_type,
            status=PolicyStatus.ACTIVE,
            sections=sections,
            raw_content=raw_content,
            variables=variables,
            references=references,
            organization_tiers=['small', 'medium', 'enterprise'],  # Default to all
            created=datetime.now().strftime('%Y-%m-%d'),
            author='Policy Committee'
        )

        return policy

    def parse_markdown(self, filepath: str) -> Policy:
        """
        Parse a Markdown file with YAML frontmatter

        Args:
            filepath: Path to Markdown file

        Returns:
            Parsed Policy object
        """
        if not FRONTMATTER_AVAILABLE:
            raise ImportError("python-frontmatter is required. Install with: pip install python-frontmatter")

        with open(filepath, 'r', encoding='utf-8') as f:
            post = frontmatter.load(f)

        metadata = post.metadata
        content = post.content

        # Build policy from metadata
        policy = Policy(
            id=metadata.get('id', self._generate_id(os.path.basename(filepath))),
            title=metadata.get('title', os.path.basename(filepath)),
            filename=os.path.basename(filepath),
            category=metadata.get('category', 'uncategorized'),
            type=PolicyType(metadata.get('type', 'policy')),
            status=PolicyStatus(metadata.get('status', 'active')),
            version=metadata.get('version', '1.0.0'),
            frameworks=metadata.get('frameworks', {}),
            references=metadata.get('references', []),
            variables=metadata.get('variables', []),
            conditions=metadata.get('conditions', []),
            requires_customization=[
                IncompletenessMarker(**m)
                for m in metadata.get('requires_customization', [])
            ],
            organization_tiers=metadata.get('organization_tiers', []),
            industries=metadata.get('industries', []),
            raw_content=content,
            created=metadata.get('created'),
            last_reviewed=metadata.get('last_reviewed'),
            next_review=metadata.get('next_review'),
            author=metadata.get('author'),
            approver=metadata.get('approver')
        )

        # Parse sections from content
        policy.sections = self._parse_markdown_sections(content)

        return policy

    def _generate_id(self, title: str) -> str:
        """Generate a kebab-case ID from title"""
        # Remove special characters and convert to kebab-case
        id_str = re.sub(r'[^\w\s-]', '', title.lower())
        id_str = re.sub(r'[\s_]+', '-', id_str)
        id_str = re.sub(r'-+', '-', id_str)
        return id_str.strip('-')

    def _detect_type(self, title: str) -> PolicyType:
        """Detect policy type from title"""
        title_lower = title.lower()

        if 'procedure' in title_lower:
            return PolicyType.PROCEDURE
        elif 'plan' in title_lower:
            return PolicyType.PLAN
        elif 'agreement' in title_lower:
            return PolicyType.AGREEMENT
        elif 'standard' in title_lower:
            return PolicyType.STANDARD
        elif 'guideline' in title_lower:
            return PolicyType.GUIDELINE
        else:
            return PolicyType.POLICY

    def _detect_category(self, title: str, doc: Document = None) -> str:
        """Detect category from title and content"""
        title_lower = title.lower()

        # Check title against category keywords
        for category, keywords in self.CATEGORY_KEYWORDS.items():
            for keyword in keywords:
                if keyword in title_lower:
                    return category

        # If doc provided, check content
        if doc:
            full_text = ' '.join([p.text.lower() for p in doc.paragraphs[:10]])
            for category, keywords in self.CATEGORY_KEYWORDS.items():
                for keyword in keywords:
                    if keyword in full_text:
                        return category

        return 'general'

    def _extract_sections(self, doc: Document) -> List[PolicySection]:
        """Extract sections from DOCX document"""
        sections = []
        current_section = None
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a heading (Roman numeral or styled)
            is_heading, level, number, title = self._parse_heading(para)

            if is_heading:
                # Save previous section
                if current_section:
                    current_section.content = '\n'.join(current_content)
                    sections.append(current_section)

                # Start new section
                current_section = PolicySection(
                    number=number,
                    title=title,
                    content='',
                    level=level
                )
                current_content = []
            else:
                current_content.append(text)

        # Save last section
        if current_section:
            current_section.content = '\n'.join(current_content)
            sections.append(current_section)

        return sections

    def _parse_heading(self, para) -> tuple:
        """
        Parse a paragraph to determine if it's a heading

        Returns:
            (is_heading, level, number, title)
        """
        text = para.text.strip()

        # Check for Roman numeral sections (I., II., etc.)
        roman_match = re.match(r'^(I{1,3}|IV|VI{0,3}|IX|X)\.\s*(.+)$', text)
        if roman_match:
            numeral = roman_match.group(1)
            title = roman_match.group(2)
            level = 2
            return (True, level, numeral, title)

        # Check for letter sections (A., B., etc.)
        letter_match = re.match(r'^([A-Z])\.\s*(.+)$', text)
        if letter_match:
            letter = letter_match.group(1)
            title = letter_match.group(2)
            level = 3
            return (True, level, letter, title)

        # Check for numbered sections (1., 2., etc.)
        number_match = re.match(r'^(\d+)\.\s*(.+)$', text)
        if number_match:
            number = number_match.group(1)
            title = number_match.group(2)
            level = 4
            return (True, level, number, title)

        # Check paragraph style for heading
        if para.style and 'heading' in para.style.name.lower():
            level = int(re.search(r'\d', para.style.name) or 2)
            return (True, level, '', text)

        return (False, 0, '', '')

    def _extract_raw_content(self, doc: Document) -> str:
        """Extract raw text content from document"""
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return '\n\n'.join(paragraphs)

    def _find_variables(self, content: str) -> List[str]:
        """Find template variables used in content"""
        found = set()

        # Look for known variable patterns
        # Pattern: ABC Company, Executive Management, etc.
        if 'ABC Company' in content:
            found.add('ORGANIZATION_NAME')
        if 'Executive Management' in content:
            found.add('EXEC_MGMT')
        if 'Chief Security Officer' in content or 'CSO' in content:
            found.add('CSO_TITLE')
        if 'IT Staff' in content:
            found.add('IT_STAFF')
        if 'Policy Committee' in content:
            found.add('APPROVER')
        if 'Risk Management Officer' in content:
            found.add('RMO_TITLE')

        # Look for date patterns
        if re.search(r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+20XX', content):
            found.add('EFFECTIVE_DATE')
            found.add('APPROVAL_DATE')

        # Look for version pattern
        if re.search(r'Version\s+\d+\.\d+', content):
            found.add('VERSION')

        return sorted(list(found))

    def _find_references(self, content: str) -> List[str]:
        """Find cross-references to other policies"""
        references = set()

        # Common reference patterns
        patterns = [
            r'(?:refer to|see|per|in accordance with|as defined in)\s+(?:the\s+)?([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+Policy)',
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+Policy)\s+(?:defines|outlines|establishes)',
            r'(?:the\s+)?([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Policy|Procedure|Plan))'
        ]

        for pattern in patterns:
            matches = re.findall(pattern, content)
            for match in matches:
                ref_id = self._generate_id(match)
                if ref_id and len(ref_id) > 5:  # Filter out very short matches
                    references.add(ref_id)

        return sorted(list(references))

    def _parse_markdown_sections(self, content: str) -> List[PolicySection]:
        """Parse sections from Markdown content"""
        sections = []
        current_section = None
        current_content = []

        for line in content.split('\n'):
            # Check for heading
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)

            if heading_match:
                # Save previous section
                if current_section:
                    current_section.content = '\n'.join(current_content)
                    sections.append(current_section)

                level = len(heading_match.group(1))
                title_text = heading_match.group(2)

                # Extract number if present
                number_match = re.match(r'^([IVX]+|[A-Z]|\d+)\.\s*(.+)$', title_text)
                if number_match:
                    number = number_match.group(1)
                    title = number_match.group(2)
                else:
                    number = ''
                    title = title_text

                current_section = PolicySection(
                    number=number,
                    title=title,
                    content='',
                    level=level
                )
                current_content = []
            else:
                current_content.append(line)

        # Save last section
        if current_section:
            current_section.content = '\n'.join(current_content)
            sections.append(current_section)

        return sections

    def load_all_policies(self, directory: str) -> Dict[str, Policy]:
        """
        Load all policies from a directory

        Args:
            directory: Path to policies directory

        Returns:
            Dictionary of policy_id -> Policy
        """
        policies = {}

        for root, dirs, files in os.walk(directory):
            for file in files:
                filepath = os.path.join(root, file)

                try:
                    if file.endswith('.docx') and not file.startswith('~'):
                        policy = self.parse_docx(filepath)
                        policies[policy.id] = policy
                    elif file.endswith('.md'):
                        policy = self.parse_markdown(filepath)
                        policies[policy.id] = policy
                except Exception as e:
                    print(f"Error parsing {filepath}: {e}")

        self._policies_cache = policies
        return policies

    def get_policy(self, policy_id: str) -> Optional[Policy]:
        """Get a policy by ID"""
        return self._policies_cache.get(policy_id)

    def save_policy(self, policy: Policy, output_dir: str) -> str:
        """
        Save a policy to Markdown + YAML format

        Args:
            policy: Policy to save
            output_dir: Output directory

        Returns:
            Path to saved file
        """
        # Create category subdirectory
        category_dir = os.path.join(output_dir, policy.category)
        os.makedirs(category_dir, exist_ok=True)

        # Generate filename
        filename = f"{policy.id}.md"
        filepath = os.path.join(category_dir, filename)

        # Write content
        content = policy.to_markdown()
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)

        return filepath


def convert_docx_library(source_dir: str, output_dir: str) -> Dict[str, str]:
    """
    Convert an entire DOCX policy library to Markdown format

    Args:
        source_dir: Directory containing DOCX files
        output_dir: Directory for Markdown output

    Returns:
        Dictionary of original_filename -> output_path
    """
    parser = PolicyParser()
    results = {}

    for file in os.listdir(source_dir):
        if file.endswith('.docx') and not file.startswith('~'):
            source_path = os.path.join(source_dir, file)

            try:
                policy = parser.parse_docx(source_path)
                output_path = parser.save_policy(policy, output_dir)
                results[file] = output_path
                print(f"✓ Converted: {file} -> {output_path}")
            except Exception as e:
                results[file] = f"ERROR: {e}"
                print(f"✗ Failed: {file} - {e}")

    return results
